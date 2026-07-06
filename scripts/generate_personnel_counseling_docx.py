#!/usr/bin/env python3
"""Generate NGC Personnel Counseling Form as Word doc for Google Drive / Google Docs."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "external_docs/templates/personnel_counseling/assets/ngc-logo.png"
OUT_DIRS = [
    ROOT / "external_docs/My Drive/Management",
    Path.home()
    / "Library/CloudStorage/GoogleDrive-neighborhoodgolfcarts985@gmail.com/My Drive/Management",
]
OUT_NAME = "NGC Personnel Counseling Form.docx"

NGC_BLUE = RGBColor(0x1A, 0x4D, 0x7A)
NGC_GREEN = RGBColor(0x6F, 0xAA, 0x2D)
MUTED = RGBColor(0x55, 0x55, 0x55)

COUNSELING_TYPES = [
    "Coaching / performance discussion",
    "Verbal warning",
    "Written warning",
    "Final written warning",
    "Policy reminder (no discipline)",
    "Positive recognition / documented coaching",
]

POLICY_AREAS = [
    "Attendance & punctuality (Mon–Fri 8–5)",
    "Shop safety & PPE",
    "Vehicle / battery safety (tow switch, disconnect, lithium handling)",
    "Diagnostic & repair procedures",
    "Job documentation (HCP notes, photos, fault codes)",
    "Quality of work / comebacks",
    "Productivity & time management",
    "Customer service & professionalism",
    "Phone / front office conduct (office staff)",
    "Deposits, parts orders & payment policy",
    "Pickup & delivery procedures (driver)",
    "Use of company tools, vehicles & property",
    "Cell phone / distraction on shop floor",
    "Teamwork & respect / harassment-free workplace",
    "Confidentiality (customer & business information)",
    "Substance use / fitness for duty",
    "Other (specify in policy field below)",
]


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = NGC_BLUE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    # underline effect via border on paragraph - skip, keep simple


def add_labeled_line(doc: Document, label: str, placeholder: str = "") -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = MUTED
    r2 = p.add_run("_" * 48 if not placeholder else f"{placeholder}{'_' * max(8, 40 - len(placeholder))}")
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)


def add_checkbox_list(doc: Document, items: list[str], cols: int = 2) -> None:
    table = doc.add_table(rows=0, cols=cols)
    table.autofit = True
    for i in range(0, len(items), cols):
        row = table.add_row()
        for c in range(cols):
            idx = i + c
            cell = row.cells[c]
            cell.text = f"☐  {items[idx]}" if idx < len(items) else ""


def add_text_area(doc: Document, label: str, lines: int = 3, hint: str = "") -> None:
    add_section_heading(doc, label)
    if hint:
        p = doc.add_paragraph(hint)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = MUTED
    for _ in range(lines):
        doc.add_paragraph("_" * 90)


def add_signature_block(doc: Document) -> None:
    table = doc.add_table(rows=4, cols=2)
    pairs = [
        ("Employee signature", "Date"),
        ("Supervisor / manager signature", "Date"),
        ("Witness signature (if applicable)", "Date"),
        ("HR / owner copy filed by", "Date filed"),
    ]
    for i, (left, right) in enumerate(pairs):
        table.rows[i].cells[0].text = f"{left}\n\n_______________________________"
        table.rows[i].cells[1].text = f"{right}\n\n_______________"


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    # Header table: logo + company info
    header = doc.add_table(rows=1, cols=2)
    header.columns[0].width = Inches(1.4)
    header.columns[1].width = Inches(5.1)
    if LOGO.exists():
        header.rows[0].cells[0].paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(1.15))
    info = header.rows[0].cells[1].paragraphs[0]
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    t = info.add_run("Neighborhood Golf Carts\n")
    t.bold = True
    t.font.size = Pt(16)
    t.font.color.rgb = NGC_BLUE
    sub = info.add_run(
        "71363 Thelma Ln, Suite E · Covington, LA 70433\n"
        "985-402-1206 · contact@ngcgolfcarts.com · NGCGolfCarts.com\n"
        "Hours: Monday–Friday, 8:00 AM – 5:00 PM"
    )
    sub.font.size = Pt(9)
    sub.font.color.rgb = MUTED

    doc.add_paragraph()

    # Title bar
    title_table = doc.add_table(rows=1, cols=1)
    title_cell = title_table.rows[0].cells[0]
    set_cell_shading(title_cell, "1A4D7A")
    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("PERSONNEL COUNSELING FORM")
    tr.bold = True
    tr.font.size = Pt(13)
    tr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    add_section_heading(doc, "Type of counseling")
    add_checkbox_list(doc, COUNSELING_TYPES)

    add_section_heading(doc, "Session information")
    fields = [
        ("Employee name", ""),
        ("Job title", "e.g. Golf cart technician"),
        ("Date of counseling", ""),
        ("Time", ""),
        ("Location", "Office / shop floor / private area"),
        ("Supervisor / manager conducting session", "Ryan White or Christine White"),
        ("Witness (optional)", ""),
        ("Prior counseling on this issue?", "Y/N — date if yes"),
    ]
    for label, ph in fields:
        add_labeled_line(doc, label, ph)

    add_section_heading(doc, "Company policy area(s) addressed")
    p = doc.add_paragraph(
        "Reference: NGC Policies & Procedures (Google Drive). Check all that apply."
    )
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = MUTED
    add_checkbox_list(doc, POLICY_AREAS)
    add_labeled_line(doc, "Specific policy or procedure cited (optional)", "")

    add_text_area(
        doc,
        "Description of concern (facts — dates, times, specific behavior)",
        4,
        "Observable facts only. Use job # instead of customer name when possible.",
    )
    add_text_area(doc, "Impact on shop, customers, team, or safety", 3)
    add_text_area(doc, "Expected behavior / corrective action", 3)
    add_text_area(doc, "Support & resources offered", 2)
    add_text_area(doc, "Consequences if issue continues", 2)
    add_text_area(doc, "Employee comments / response", 3)

    add_section_heading(doc, "Acknowledgment")
    ack = doc.add_paragraph(
        "I acknowledge that I received counseling regarding the matter described above. "
        "My signature does not necessarily indicate agreement with the contents, only that "
        "the discussion occurred and I received a copy of this form."
    )
    ack.runs[0].font.size = Pt(10)
    add_signature_block(doc)

    doc.add_paragraph()
    notice = doc.add_paragraph(
        "CONFIDENTIAL PERSONNEL RECORD. Store completed forms in the secure Management folder "
        "(Google Drive). Do not store in customer-facing systems. Neighborhood Golf Carts is an "
        "at-will employer subject to applicable federal and Louisiana law. This form is a "
        "management tool and does not modify at-will status unless separately agreed in writing. "
        "Consult qualified employment counsel for termination or complex situations."
    )
    notice.runs[0].font.size = Pt(8)
    notice.runs[0].font.color.rgb = MUTED

    footer = doc.add_paragraph("NGC Personnel Counseling Form · Rev. 2026-07          Form ID: HR-COUNSEL-001")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED

    return doc


def main() -> int:
    doc = build_document()
    written = []
    for out_dir in OUT_DIRS:
        if not out_dir.parent.exists() and "CloudStorage" in str(out_dir):
            continue
        out_dir.mkdir(parents=True, exist_ok=True)
        path = out_dir / OUT_NAME
        doc.save(path)
        written.append(path)
    if not written:
        fallback = ROOT / "external_docs/My Drive/Management"
        fallback.mkdir(parents=True, exist_ok=True)
        path = fallback / OUT_NAME
        doc.save(path)
        written.append(path)
    for p in written:
        print(f"Wrote {p}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
